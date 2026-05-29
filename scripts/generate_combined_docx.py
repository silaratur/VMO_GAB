#!/usr/bin/env python3
"""
Gera documento Word único consolidado (.docx) para o pacote DEM-2026-008.
Layout profissional com capa, sumário, separadores de fase e todos os documentos.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import lxml.etree as etree

# ── Paleta de cores ────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1E, 0x3C, 0x78)
BLUE_MID   = RGBColor(0x1E, 0x50, 0xA0)
BLUE_LIGHT = RGBColor(0x32, 0x66, 0xCC)
GREY_TEXT  = RGBColor(0x55, 0x55, 0x55)
GREY_LIGHT = RGBColor(0xF5, 0xF8, 0xFC)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GOLD       = RGBColor(0xF0, 0xC0, 0x40)
TABLE_HEAD = RGBColor(0xDC, 0xE6, 0xF5)
TABLE_ALT  = RGBColor(0xF5, 0xF8, 0xFC)
RED_ALERT  = RGBColor(0xC0, 0x20, 0x20)
GREEN_OK   = RGBColor(0x1A, 0x7A, 0x3C)

# ── Ordem dos documentos ──────────────────────────────────────────────────────
DOC_ORDER = [
    ("01-qualificacao", "demanda-coletada.md",     "Demanda Coletada"),
    ("01-qualificacao", "gate-intake.md",           "Gate de Governança — Intake"),
    ("01-qualificacao", "qualificacao.md",          "Qualificação da Demanda"),
    ("01-qualificacao", "gate-qualificacao.md",     "Gate de Governança — Qualificação"),
    ("01-qualificacao", "qualificacao-aprovada.md", "Qualificação Aprovada"),
    ("02-iniciacao",    "documentacao-base.md",     "Documentação Base (TAP + PM Canvas)"),
    ("02-iniciacao",    "requisitos.md",            "Especificação de Requisitos (ERF)"),
    ("02-iniciacao",    "work-request.md",          "Work Request — Mini-RFP"),
    ("03-planejamento", "cronograma.md",            "Cronograma do Projeto"),
    ("03-planejamento", "plano-riscos.md",          "Plano de Riscos"),
    ("03-planejamento", "kpis.md",                  "Framework de KPIs"),
    ("04-monitoramento","status-report-2026-05-28.md","Status Report Inicial"),
    ("05-encerramento", "revisao-final.md",         "Revisão de Qualidade"),
    ("05-encerramento", "auditoria-governanca.md",  "Auditoria de Governança"),
    ("05-encerramento", "aprovacao-final.md",       "Aprovação Final"),
]

SECTION_LABELS = {
    "01-qualificacao":  ("FASE 01", "QUALIFICAÇÃO"),
    "02-iniciacao":     ("FASE 02", "INICIAÇÃO"),
    "03-planejamento":  ("FASE 03", "PLANEJAMENTO"),
    "04-monitoramento": ("FASE 04", "MONITORAMENTO"),
    "05-encerramento":  ("FASE 05", "ENCERRAMENTO"),
}

SECTION_ICONS = {
    "01-qualificacao":  "🔍",
    "02-iniciacao":     "📋",
    "03-planejamento":  "📅",
    "04-monitoramento": "📊",
    "05-encerramento":  "✅",
}


# ── Helpers XML ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Define cor de fundo de célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace('#', ''))
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 6)))
            el.set(qn('w:color'), val.get('color', '1E3C78'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br


def insert_page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def set_paragraph_border_bottom(para, color='1E3C78', sz=12):
    """Linha abaixo do parágrafo."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_run_highlight(run, color_hex):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    rPr.append(shd)


def add_bookmark(para, name):
    """Adiciona marcador para sumário interno."""
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(abs(hash(name)) % 10000))
    bm_start.set(qn('w:name'), name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(abs(hash(name)) % 10000))
    para._p.insert(0, bm_start)
    para._p.append(bm_end)


def set_doc_margins(section, top=2.5, bottom=2.0, left=2.5, right=2.0):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


# ── Estilos de parágrafo ──────────────────────────────────────────────────────

def style_normal(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15


def apply_heading_style(para, level, doc):
    """Aplica formatação de heading customizado."""
    font = para.runs[0].font if para.runs else para.add_run().font
    para.paragraph_format.space_before = Pt([14, 10, 7, 5][min(level - 1, 3)])
    para.paragraph_format.space_after = Pt([8, 6, 4, 3][min(level - 1, 3)])
    para.paragraph_format.keep_with_next = True

    if level == 1:
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = BLUE_DARK
        set_paragraph_border_bottom(para, '1E3C78', 18)

    elif level == 2:
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = BLUE_MID
        set_paragraph_border_bottom(para, '1E50A0', 10)

    elif level == 3:
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = BLUE_LIGHT

    elif level == 4:
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = BLUE_LIGHT


# ── Capa ──────────────────────────────────────────────────────────────────────

def add_cover(doc):
    section = doc.sections[0]
    set_doc_margins(section)

    # Faixa superior azul (simulada com tabela)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_bg(cell, '1E3C78')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('VMO AUTÔNOMO')
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = WHITE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('PACOTE DE INICIAÇÃO DE PROJETO')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0xC8, 0xD8, 0xF0)
    p2.paragraph_format.space_after = Pt(16)

    doc.add_paragraph()

    # ID do projeto em destaque
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c2 = tbl2.rows[0].cells[0]
    set_cell_bg(c2, 'F0C040')
    set_cell_borders(c2,
        top={'val': 'single', 'sz': 12, 'color': '1E3C78'},
        bottom={'val': 'single', 'sz': 12, 'color': '1E3C78'})
    p3 = c2.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(10)
    p3.paragraph_format.space_after = Pt(10)
    r3 = p3.add_run('DEM-2026-008')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(22)
    r3.font.bold = True
    r3.font.color.rgb = BLUE_DARK

    doc.add_paragraph()

    # Título da demanda
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('Integração SGMM03 — Campos Empresa e Contrato (InterCompany)')
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(14)
    r_title.font.bold = True
    r_title.font.color.rgb = BLUE_DARK
    p_title.paragraph_format.space_after = Pt(4)

    p_ch = doc.add_paragraph()
    p_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ch = p_ch.add_run('Chamado #6800446')
    r_ch.font.name = 'Calibri'
    r_ch.font.size = Pt(12)
    r_ch.font.color.rgb = GREY_TEXT
    p_ch.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # Tabela de metadados
    meta = [
        ("Sistema",         "Sistemática ERP — Solicitação de Novas Demandas / Projetos"),
        ("Solicitante",     "Jenifer dos Santos Carvalho — VIX Matriz"),
        ("Responsável",     "Mara Rubia Silva Rocha — Holding DTI"),
        ("Grupo Sol.",      "Projetos DTI / Holding DTI"),
        ("Data Abertura",   "08/05/2026"),
        ("Prazo Original",  "15/05/2026 (SLA em atraso — 81h42)"),
        ("Data Pacote",     "28/05/2026"),
        ("Status",          "APROVADO COM CONDIÇÕES — Score 50/100"),
        ("Classificação",   "Melhoria Evolutiva — Time SAP PM/FI"),
        ("Revisão VMO",     "8,68/10 (Vera Veredito) | Auditoria: APROVADO COM RESSALVAS"),
    ]
    tbl3 = doc.add_table(rows=len(meta), cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        row = tbl3.rows[i]
        c_k = row.cells[0]
        c_v = row.cells[1]
        c_k.width = Cm(4.5)
        c_v.width = Cm(12)
        if i % 2 == 0:
            set_cell_bg(c_k, 'DCE6F5')
            set_cell_bg(c_v, 'F5F8FC')
        else:
            set_cell_bg(c_k, 'C8D8F0')
            set_cell_bg(c_v, 'FFFFFF')
        pk = c_k.paragraphs[0]
        pk.paragraph_format.space_before = Pt(4)
        pk.paragraph_format.space_after = Pt(4)
        rk = pk.add_run(k)
        rk.font.name = 'Calibri'
        rk.font.size = Pt(9)
        rk.font.bold = True
        rk.font.color.rgb = BLUE_DARK
        pv = c_v.paragraphs[0]
        pv.paragraph_format.space_before = Pt(4)
        pv.paragraph_format.space_after = Pt(4)
        rv = pv.add_run(v)
        rv.font.name = 'Calibri'
        rv.font.size = Pt(9)
        if k == "Status":
            rv.font.bold = True
            rv.font.color.rgb = GREEN_OK

    doc.add_paragraph()

    # Rodapé da capa
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_conf = p_conf.add_run('CONFIDENCIAL — VMO Autônomo / Holding DTI — 2026')
    r_conf.font.name = 'Calibri'
    r_conf.font.size = Pt(8)
    r_conf.font.italic = True
    r_conf.font.color.rgb = GREY_TEXT


# ── Sumário ───────────────────────────────────────────────────────────────────

def add_toc(doc):
    insert_page_break(doc)

    p = doc.add_paragraph()
    r = p.add_run('SUMÁRIO')
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BLUE_DARK
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_border_bottom(p, '1E3C78', 18)

    doc.add_paragraph()

    current_fase = None
    for fase, fname, title in DOC_ORDER:
        if fase != current_fase:
            current_fase = fase
            fase_num, fase_nome = SECTION_LABELS[fase]
            icon = SECTION_ICONS[fase]
            p_fase = doc.add_paragraph()
            p_fase.paragraph_format.space_before = Pt(8)
            p_fase.paragraph_format.space_after = Pt(2)
            r_fase = p_fase.add_run(f'{icon}  {fase_num} — {fase_nome}')
            r_fase.font.name = 'Calibri'
            r_fase.font.size = Pt(11)
            r_fase.font.bold = True
            r_fase.font.color.rgb = BLUE_DARK

        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.left_indent = Cm(0.8)
        p_item.paragraph_format.space_before = Pt(0)
        p_item.paragraph_format.space_after = Pt(1)
        r_item = p_item.add_run(title)
        r_item.font.name = 'Calibri'
        r_item.font.size = Pt(10)
        r_item.font.color.rgb = BLUE_MID

    doc.add_paragraph()

    # Nota
    p_note = doc.add_paragraph()
    r_note = p_note.add_run('Documento gerado automaticamente pelo VMO Autônomo em 28/05/2026. '
                             'Contém 15 documentos do pacote completo de iniciação do projeto DEM-2026-008.')
    r_note.font.name = 'Calibri'
    r_note.font.size = Pt(8)
    r_note.font.italic = True
    r_note.font.color.rgb = GREY_TEXT
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Separador de fase ─────────────────────────────────────────────────────────

def add_section_divider(doc, fase):
    insert_page_break(doc)
    fase_num, fase_nome = SECTION_LABELS[fase]
    icon = SECTION_ICONS[fase]

    # Bloco azul de fase
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '1E3C78')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f'{icon}  {fase_num}')
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xC8, 0xD8, 0xF0)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(14)
    p2.paragraph_format.left_indent = Cm(0.5)
    r2 = p2.add_run(fase_nome)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(22)
    r2.font.bold = True
    r2.font.color.rgb = WHITE

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Separador de documento ────────────────────────────────────────────────────

def add_doc_header(doc, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'EEF3FB')
    set_cell_borders(cell,
        top={'val': 'single', 'sz': 18, 'color': '1E50A0'},
        bottom={'val': 'single', 'sz': 6, 'color': '3266CC'})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'📄  {title}')
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = BLUE_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Parser de markdown → docx ─────────────────────────────────────────────────

def apply_inline(run_parent, text):
    """Aplica formatação inline bold/italic/code num parágrafo."""
    # Tokeniza bold, italic, code
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*'
        r'|\*\*(.+?)\*\*'
        r'|\*(.+?)\*'
        r'|`(.+?)`'
        r'|~~(.+?)~~'
        r'|\[([^\]]+)\]\([^\)]+\))'
    )
    last = 0
    for m in pattern.finditer(text):
        # texto antes
        if m.start() > last:
            r = run_parent.add_run(text[last:m.start()])
            r.font.name = 'Calibri'

        full = m.group(0)
        if full.startswith('***'):
            r = run_parent.add_run(m.group(2))
            r.font.bold = True
            r.font.italic = True
        elif full.startswith('**'):
            r = run_parent.add_run(m.group(3))
            r.font.bold = True
        elif full.startswith('*'):
            r = run_parent.add_run(m.group(4))
            r.font.italic = True
        elif full.startswith('`'):
            r = run_parent.add_run(m.group(5))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            set_run_highlight(r, 'F0F0F0')
        elif full.startswith('~~'):
            r = run_parent.add_run(m.group(6))
            r.font.strike = True
        elif full.startswith('['):
            r = run_parent.add_run(m.group(7))
            r.font.underline = True
            r.font.color.rgb = BLUE_LIGHT

        if r:
            r.font.name = r.font.name or 'Calibri'
        last = m.end()

    if last < len(text):
        r = run_parent.add_run(text[last:])
        r.font.name = 'Calibri'


def add_md_table(doc, table_buf):
    """Renderiza tabela markdown no docx."""
    rows_data = []
    header_done = False
    is_header_row = True
    for raw_row in table_buf:
        cells = [c.strip() for c in raw_row.strip().strip('|').split('|')]
        if all(re.match(r'^[-: ]+$', c) for c in cells if c):
            header_done = True
            is_header_row = False
            continue
        rows_data.append((cells, is_header_row and not header_done))
        is_header_row = False

    if not rows_data:
        return

    ncols = max(len(r[0]) for r in rows_data)
    for r in rows_data:
        while len(r[0]) < ncols:
            r[0].append('')

    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    available_width = Cm(16.5)
    col_w = available_width / ncols
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_w

    for ri, (cells, is_hdr) in enumerate(rows_data):
        row = tbl.rows[ri]
        for ci, text in enumerate(cells[:ncols]):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if is_hdr:
                set_cell_bg(cell, 'DCE6F5')
            elif ri % 2 == 0:
                set_cell_bg(cell, 'FFFFFF')
            else:
                set_cell_bg(cell, 'F5F8FC')

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.1)

            clean = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
            apply_inline(p, clean)

            for run in p.runs:
                run.font.name = run.font.name or 'Calibri'
                run.font.size = Pt(8.5)
                if is_hdr:
                    run.font.bold = True
                    run.font.color.rgb = BLUE_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def parse_md_to_doc(doc, md_path):
    """Converte markdown para elementos docx."""
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_md_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        raw = lines[i].rstrip('\n')

        # ── bloco de código ─────────────────────────────────────────────────
        if raw.strip().startswith('```'):
            if in_code:
                if code_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.right_indent = Cm(0.5)
                    r = p.add_run('\n'.join(code_buf))
                    r.font.name = 'Courier New'
                    r.font.size = Pt(8.5)
                    set_run_highlight(r, 'F0F0F0')
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── tabela ──────────────────────────────────────────────────────────
        if raw.strip().startswith('|'):
            table_buf.append(raw)
            i += 1
            continue
        else:
            flush_table()

        # ── headings ────────────────────────────────────────────────────────
        if raw.startswith('# ') and not raw.startswith('## '):
            p = doc.add_paragraph()
            apply_inline(p, raw[2:])
            apply_heading_style(p, 1, doc)

        elif raw.startswith('## '):
            p = doc.add_paragraph()
            apply_inline(p, raw[3:])
            apply_heading_style(p, 2, doc)

        elif raw.startswith('### '):
            p = doc.add_paragraph()
            apply_inline(p, raw[4:])
            apply_heading_style(p, 3, doc)

        elif raw.startswith('#### '):
            p = doc.add_paragraph()
            apply_inline(p, raw[5:])
            apply_heading_style(p, 4, doc)

        # ── HR ──────────────────────────────────────────────────────────────
        elif raw.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            set_paragraph_border_bottom(p, 'CCCCCC', 6)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

        # ── blockquote ──────────────────────────────────────────────────────
        elif raw.strip().startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            set_paragraph_border_bottom(p, 'CCCCCC', 4)
            apply_inline(p, raw.strip()[2:])
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = GREY_TEXT
                run.font.size = Pt(9.5)

        # ── checkbox ────────────────────────────────────────────────────────
        elif re.match(r'^\s*- \[[ xX]\]', raw):
            checked = bool(re.match(r'^\s*- \[[xX]\]', raw))
            text = re.sub(r'^\s*- \[.\]\s*', '', raw)
            indent = len(raw) - len(raw.lstrip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4 + indent * 0.02)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            mark = '☑' if checked else '☐'
            r_mark = p.add_run(f'{mark}  ')
            r_mark.font.name = 'Calibri'
            r_mark.font.size = Pt(9.5)
            r_mark.font.color.rgb = BLUE_MID if checked else GREY_TEXT
            apply_inline(p, text)
            for run in p.runs[1:]:
                run.font.size = Pt(9.5)

        # ── lista com marcadores ─────────────────────────────────────────────
        elif re.match(r'^\s*[-*+] ', raw):
            indent = len(raw) - len(raw.lstrip())
            text = re.sub(r'^\s*[-*+] ', '', raw)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.025)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r_bullet = p.add_run('•  ')
            r_bullet.font.name = 'Calibri'
            r_bullet.font.size = Pt(10)
            r_bullet.font.color.rgb = BLUE_MID
            apply_inline(p, text)
            for run in p.runs[1:]:
                run.font.name = run.font.name or 'Calibri'
                run.font.size = Pt(10)

        # ── lista numerada ───────────────────────────────────────────────────
        elif re.match(r'^\s*\d+\.\s', raw):
            indent = len(raw) - len(raw.lstrip())
            num = re.match(r'^\s*(\d+)\.', raw).group(1)
            text = re.sub(r'^\s*\d+\.\s+', '', raw)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.025)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r_num = p.add_run(f'{num}.  ')
            r_num.font.name = 'Calibri'
            r_num.font.size = Pt(10)
            r_num.font.bold = True
            r_num.font.color.rgb = BLUE_MID
            apply_inline(p, text)
            for run in p.runs[1:]:
                run.font.name = run.font.name or 'Calibri'
                run.font.size = Pt(10)

        # ── linha vazia ──────────────────────────────────────────────────────
        elif raw.strip() == '':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

        # ── parágrafo normal ─────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            apply_inline(p, raw)
            for run in p.runs:
                run.font.name = run.font.name or 'Calibri'
                run.font.size = Pt(10)

        i += 1

    flush_table()
    if in_code and code_buf:
        p = doc.add_paragraph()
        r = p.add_run('\n'.join(code_buf))
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        set_run_highlight(r, 'F0F0F0')


# ── Header/Footer do documento ────────────────────────────────────────────────

def add_header_footer(doc):
    """Adiciona cabeçalho e rodapé a todas as seções após a capa."""
    for i, section in enumerate(doc.sections):
        if i == 0:
            # Capa sem header/footer
            section.different_first_page_header_footer = True
            continue

        # Header
        header = section.header
        header.is_linked_to_previous = False
        htbl = header.add_table(rows=1, cols=3, width=Cm(17))
        htbl.style = 'Table Grid'
        l_cell, m_cell, r_cell = htbl.rows[0].cells
        for cell in [l_cell, m_cell, r_cell]:
            set_cell_bg(cell, 'EEF3FB')

        p_l = l_cell.paragraphs[0]
        r_l = p_l.add_run('VMO Autônomo — DEM-2026-008')
        r_l.font.name = 'Calibri'
        r_l.font.size = Pt(8)
        r_l.font.bold = True
        r_l.font.color.rgb = BLUE_DARK

        p_m = m_cell.paragraphs[0]
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_m = p_m.add_run('Integração SGMM03 — InterCompany')
        r_m.font.name = 'Calibri'
        r_m.font.size = Pt(8)
        r_m.font.color.rgb = GREY_TEXT

        p_r = r_cell.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_r = p_r.add_run('Chamado #6800446')
        r_r.font.name = 'Calibri'
        r_r.font.size = Pt(8)
        r_r.font.color.rgb = GREY_TEXT

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        ftbl = footer.add_table(rows=1, cols=3, width=Cm(17))
        fl, fm, fr = ftbl.rows[0].cells

        p_fl = fl.paragraphs[0]
        r_fl = p_fl.add_run('CONFIDENCIAL — Holding DTI')
        r_fl.font.name = 'Calibri'
        r_fl.font.size = Pt(7.5)
        r_fl.font.color.rgb = GREY_TEXT

        p_fm = fm.paragraphs[0]
        p_fm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Número de página automático
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_pn = p_fm.add_run()
        run_pn._r.append(fldChar1)
        run_pn._r.append(instrText)
        run_pn._r.append(fldChar2)
        run_pn.font.name = 'Calibri'
        run_pn.font.size = Pt(8)
        run_pn.font.color.rgb = GREY_TEXT

        p_fr = fr.paragraphs[0]
        p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_fr = p_fr.add_run('28/05/2026')
        r_fr.font.name = 'Calibri'
        r_fr.font.size = Pt(7.5)
        r_fr.font.color.rgb = GREY_TEXT


# ── Main ──────────────────────────────────────────────────────────────────────

def generate(project_dir, output_path):
    doc = Document()
    style_normal(doc)

    # Margens globais
    for section in doc.sections:
        set_doc_margins(section)

    print('\nGerando DOCX consolidado...\n')

    # Capa
    add_cover(doc)
    insert_page_break(doc)

    # Sumário
    add_toc(doc)

    # Documentos por fase
    current_fase = None
    for fase, fname, title in DOC_ORDER:
        md_path = os.path.join(project_dir, fase, fname)
        if not os.path.exists(md_path):
            print(f'  AVISO: não encontrado — {md_path}')
            continue

        if fase != current_fase:
            add_section_divider(doc, fase)
            current_fase = fase

        add_doc_header(doc, title)
        parse_md_to_doc(doc, md_path)
        insert_page_break(doc)
        print(f'  ✓  {fase}/{fname}')

    # Header/Footer
    add_header_footer(doc)

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f'\n✅  DOCX gerado: {output_path} ({size_kb} KB)\n')


if __name__ == '__main__':
    project_dir = '/home/user/VMO_GAB/squads/vmo-autonomo/projects/DEM-2026-008'
    output_path = '/home/user/VMO_GAB/squads/vmo-autonomo/projects/DEM-2026-008/DEM-2026-008-pacote-completo.docx'
    generate(project_dir, output_path)
