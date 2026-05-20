#!/usr/bin/env python3
"""
Convert Markdown files to PDF and/or DOCX.

Usage:
  python3 md_to_docs.py <file.md>                    # PDF (default)
  python3 md_to_docs.py <file.md> --format pdf
  python3 md_to_docs.py <file.md> --format docx
  python3 md_to_docs.py <file.md> --format both

PDF output: formal project document layout (header, page numbers, footer,
professional typography and table styling).
"""

import sys
import re
import argparse
from pathlib import Path


# ── PDF ───────────────────────────────────────────────────────────────────────

def md_to_pdf(md_path: Path) -> Path:
    import markdown as md_lib
    from weasyprint import HTML, CSS

    text = md_path.read_text(encoding="utf-8")

    # Extract title (first H1) and subtitle (second line after H1 if not heading)
    lines = text.split("\n")
    doc_title = ""
    doc_subtitle = ""
    doc_meta = ""
    for i, line in enumerate(lines):
        if line.startswith("# ") and not doc_title:
            doc_title = line[2:].strip()
            # Look for subtitle/meta on next non-empty lines
            for j in range(i + 1, min(i + 4, len(lines))):
                l = lines[j].strip()
                if l and not l.startswith("#") and not l.startswith("---"):
                    if not doc_subtitle:
                        doc_subtitle = l
                    elif not doc_meta:
                        doc_meta = l
                    else:
                        break
            break

    extensions = ["tables", "fenced_code", "nl2br", "attr_list"]
    html_body = md_lib.markdown(text, extensions=extensions)

    # Build formal header block
    header_html = f"""
    <div class="doc-header">
        <div class="doc-header-left">
            <div class="doc-org">VMO — Escritório de Projetos</div>
            <div class="doc-title">{doc_title}</div>
            {"<div class='doc-subtitle'>" + doc_subtitle + "</div>" if doc_subtitle else ""}
            {"<div class='doc-meta'>" + doc_meta + "</div>" if doc_meta else ""}
        </div>
        <div class="doc-header-right">
            <div class="doc-badge">DOCUMENTO FORMAL</div>
        </div>
    </div>
    <div class="doc-divider"></div>
    """

    css = CSS(string="""
        /* ── Page setup ── */
        @page {
            size: A4;
            margin: 2.2cm 2.5cm 2.8cm 2.5cm;

            @top-left {
                content: "VMO — Escritório de Projetos";
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 8pt;
                color: #7a7a7a;
                border-bottom: 0.5pt solid #cccccc;
                padding-bottom: 4pt;
            }
            @top-right {
                content: string(doc-title-running);
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 8pt;
                color: #7a7a7a;
                border-bottom: 0.5pt solid #cccccc;
                padding-bottom: 4pt;
            }
            @bottom-left {
                content: "Documento gerado pelo VMO Autônomo";
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 7.5pt;
                color: #aaaaaa;
                border-top: 0.5pt solid #e0e0e0;
                padding-top: 4pt;
            }
            @bottom-right {
                content: "Pág. " counter(page) " / " counter(pages);
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 7.5pt;
                color: #aaaaaa;
                border-top: 0.5pt solid #e0e0e0;
                padding-top: 4pt;
            }
        }

        /* ── Base ── */
        body {
            font-family: 'DejaVu Sans', Arial, sans-serif;
            font-size: 10.5pt;
            line-height: 1.55;
            color: #1a1a1a;
            background: #ffffff;
        }

        /* ── Document header block (first page) ── */
        .doc-header {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            margin-bottom: 6pt;
            padding-bottom: 10pt;
        }
        .doc-header-left { flex: 1; }
        .doc-org {
            font-size: 8pt;
            color: #7a7a7a;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 4pt;
        }
        .doc-title {
            font-size: 18pt;
            font-weight: bold;
            color: #1a3a5c;
            line-height: 1.2;
            margin-bottom: 4pt;
            string-set: doc-title-running content();
        }
        .doc-subtitle {
            font-size: 10pt;
            color: #444;
            margin-bottom: 2pt;
        }
        .doc-meta {
            font-size: 9pt;
            color: #777;
        }
        .doc-header-right {
            text-align: right;
            padding-left: 16pt;
        }
        .doc-badge {
            font-size: 7pt;
            font-weight: bold;
            color: #ffffff;
            background: #1a3a5c;
            padding: 3pt 8pt;
            border-radius: 2pt;
            text-transform: uppercase;
            letter-spacing: 0.08em;
        }
        .doc-divider {
            border: none;
            border-top: 2.5pt solid #1a3a5c;
            margin-bottom: 18pt;
        }

        /* ── Headings ── */
        h1 {
            font-size: 15pt;
            font-weight: bold;
            color: #1a3a5c;
            border-bottom: 1.5pt solid #1a3a5c;
            padding-bottom: 3pt;
            margin-top: 20pt;
            margin-bottom: 8pt;
            page-break-after: avoid;
        }
        h2 {
            font-size: 12.5pt;
            font-weight: bold;
            color: #1e5296;
            border-bottom: 0.75pt solid #1e5296;
            padding-bottom: 2pt;
            margin-top: 16pt;
            margin-bottom: 6pt;
            page-break-after: avoid;
        }
        h3 {
            font-size: 11pt;
            font-weight: bold;
            color: #2a4a7f;
            margin-top: 12pt;
            margin-bottom: 4pt;
            page-break-after: avoid;
        }
        h4 {
            font-size: 10.5pt;
            font-weight: bold;
            font-style: italic;
            color: #444444;
            margin-top: 10pt;
            margin-bottom: 3pt;
            page-break-after: avoid;
        }

        /* ── Paragraphs ── */
        p { margin: 0 0 7pt 0; }

        /* ── Tables ── */
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 10pt 0 14pt 0;
            font-size: 9.5pt;
            page-break-inside: auto;
        }
        thead { display: table-header-group; }
        th {
            background-color: #1a3a5c;
            color: #ffffff;
            font-weight: bold;
            padding: 6pt 8pt;
            border: 1pt solid #1a3a5c;
            text-align: left;
            vertical-align: top;
        }
        td {
            padding: 5pt 8pt;
            border: 0.5pt solid #c8d4e3;
            vertical-align: top;
        }
        tr:nth-child(even) td { background-color: #f0f5fb; }
        tr:nth-child(odd) td { background-color: #ffffff; }
        tr:hover td { background-color: #e4edf8; }

        /* ── Lists ── */
        ul, ol {
            margin: 4pt 0 8pt 18pt;
            padding: 0;
        }
        li { margin-bottom: 3pt; }
        li > ul, li > ol { margin-top: 2pt; }

        /* ── Code ── */
        code {
            font-family: 'DejaVu Sans Mono', 'Courier New', monospace;
            font-size: 8.5pt;
            background: #f2f4f8;
            padding: 1pt 4pt;
            border-radius: 2pt;
            color: #1a1a1a;
        }
        pre {
            background: #f2f4f8;
            border-left: 3pt solid #1a3a5c;
            padding: 8pt 12pt;
            font-size: 8pt;
            overflow-x: auto;
            margin: 8pt 0;
        }
        pre code { background: none; padding: 0; }

        /* ── Blockquote / callouts ── */
        blockquote {
            border-left: 3pt solid #1e5296;
            margin: 8pt 0;
            padding: 6pt 14pt;
            background: #f5f8fd;
            color: #333;
            font-size: 10pt;
        }
        blockquote p { margin: 0; }

        /* ── Horizontal rule ── */
        hr {
            border: none;
            border-top: 0.75pt solid #cccccc;
            margin: 14pt 0;
        }

        /* ── Strong / em ── */
        strong { color: #0d2840; font-weight: bold; }
        em { color: #333; }

        /* ── Page breaks ── */
        .page-break { page-break-after: always; }
        h1 { page-break-before: auto; }
    """)

    html_full = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <title>{doc_title}</title>
</head>
<body>
{header_html}
{html_body}
</body>
</html>"""

    out = md_path.with_suffix(".pdf")
    HTML(string=html_full, base_url=str(md_path.parent)).write_pdf(
        str(out), stylesheets=[css], presentational_hints=True
    )
    return out


# ── DOCX ──────────────────────────────────────────────────────────────────────

def md_to_docx(md_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DARK_BLUE = RGBColor(0x1A, 0x3A, 0x5C)
    MID_BLUE  = RGBColor(0x1E, 0x52, 0x96)
    LIGHT_BLUE_HEX = "D6E4F0"
    HEADER_BG_HEX  = "1A3A5C"

    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def set_run_font(run, bold=False, italic=False, size=None, color=None, name="Calibri"):
        run.font.name = name
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def add_heading(content, level):
        p = doc.add_heading("", level=level)
        run = p.add_run(content)
        if level == 1:
            set_run_font(run, bold=True, size=15, color=DARK_BLUE)
        elif level == 2:
            set_run_font(run, bold=True, size=12.5, color=MID_BLUE)
        elif level == 3:
            set_run_font(run, bold=True, size=11, color=RGBColor(0x2A, 0x4A, 0x7F))
        else:
            set_run_font(run, bold=True, italic=True, size=10.5, color=RGBColor(0x44, 0x44, 0x44))

    def add_inline(para, line):
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|_[^_]+_)', line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = para.add_run(part[2:-2])
                r.bold = True
            elif part.startswith("`") and part.endswith("`"):
                r = para.add_run(part[1:-1])
                r.font.name = "Courier New"
                r.font.size = Pt(9)
            elif part.startswith("_") and part.endswith("_"):
                r = para.add_run(part[1:-1])
                r.italic = True
            else:
                para.add_run(part)

    def flush_table(rows):
        data_rows = [r for r in rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
        if not data_rows:
            return
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            parsed.append(cells)
        col_count = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=col_count)
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(parsed):
            for c_idx in range(col_count):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                clean = re.sub(r'`([^`]+)`', r'\1', clean)
                cell = tbl.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                run = p.add_run(clean)
                run.font.size = Pt(9.5)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), HEADER_BG_HEX)
                    tcPr.append(shd)
                elif r_idx % 2 == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F0F5FB")
                    tcPr.append(shd)

    table_rows = []
    in_code_block = False
    lines_list = text.split("\n")

    for line in lines_list:
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        if line.strip().startswith("|"):
            table_rows.append(line)
            continue
        else:
            if table_rows:
                flush_table(table_rows)
                table_rows = []

        if line.startswith("# "):
            add_heading(line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(line[4:].strip(), 3)
        elif line.startswith("#### "):
            add_heading(line[5:].strip(), 4)
        elif re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            add_inline(p, line[2:].strip())
            p.paragraph_format.left_indent = Cm(0.8)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                run.font.size = Pt(10)
        elif re.match(r'^[-*•] ', line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip())
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r'^\d+\. ', '', line))
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            add_inline(p, line)

    if table_rows:
        flush_table(table_rows)

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to PDF/DOCX")
    parser.add_argument("input", help="Path to .md file")
    parser.add_argument("--format", choices=["pdf", "docx", "both"], default="pdf")
    args = parser.parse_args()

    md_path = Path(args.input).resolve()
    if not md_path.exists():
        print(f"ERROR: file not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    if args.format in ("pdf", "both"):
        out = md_to_pdf(md_path)
        print(f"PDF:  {out}")

    if args.format in ("docx", "both"):
        out = md_to_docx(md_path)
        print(f"DOCX: {out}")


if __name__ == "__main__":
    main()
